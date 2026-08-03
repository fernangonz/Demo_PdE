# -*- coding: utf-8 -*-
"""Fichas documentales desde Word en ``Fichas/*.docx`` (un archivo = un modelo).

El nombre del archivo (sin extension) se empareja con el catalogo igual que
antes las hojas de Excel: p.ej. ``PI FALTA DE FRANCOBORDO.docx`` ->
``PI Falta de francobordo`` / ``falta_francobordo_elo``.

Ruta robusta: tablas e imagenes con ``python-docx`` (+ zip OOXML).
Mammoth queda como fallback opcional si esta instalado.
"""
from __future__ import annotations

import sys

import base64
import hashlib
import html as html_lib
import logging
import re
import unicodedata
import zipfile
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path
from xml.etree import ElementTree as ET

RAIZ_PROYECTO = Path(__file__).resolve().parents[2]
CARPETA_FICHAS = RAIZ_PROYECTO / "Fichas"
CARPETA_MEDIA = CARPETA_FICHAS / "_media"

_EXTENSIONES_WORD = (".docx",)
_log = logging.getLogger(__name__)

_CSS_FICHA = (
    ".pde-ficha-word{"
    "font-family:Calibri,'Segoe UI',Arial,sans-serif;"
    "font-size:15px;line-height:1.45;color:#1a1a1a;"
    "}"
    ".pde-ficha-word table{"
    "border-collapse:collapse;width:100%;margin:0.75em 0;"
    "}"
    ".pde-ficha-word th,.pde-ficha-word td{"
    "border:1px solid #333;padding:6px 8px;vertical-align:top;"
    "}"
    ".pde-ficha-word p{margin:0.4em 0;}"
    ".pde-ficha-word img{max-width:100%;height:auto;display:block;margin:8px auto;}"
    ".pde-ficha-word h1,.pde-ficha-word h2,.pde-ficha-word h3{"
    "margin:0.8em 0 0.35em;font-weight:700;"
    "}"
    ".pde-ficha-eq{"
    "margin:0.75em 0 0.25em;padding:0.55em 0.7em;"
    "background:#f7f7f7;border-left:3px solid #333;"
    "font-family:'Cambria Math',Cambria,'Times New Roman',serif;"
    "font-size:1.05em;line-height:1.55;overflow-x:auto;"
    "}"
    ".pde-ficha-eq .frac{"
    "display:inline-block;vertical-align:middle;text-align:center;margin:0 0.15em;"
    "}"
    ".pde-ficha-eq .frac .num{"
    "display:block;border-bottom:1px solid #333;padding:0 0.35em 0.1em;"
    "line-height:1.2;"
    "}"
    ".pde-ficha-eq .frac .den{"
    "display:block;padding:0.1em 0.35em 0;line-height:1.2;"
    "}"
    ".pde-ficha-eq sub{font-size:0.75em;}"
)


@dataclass(frozen=True)
class EcuacionFicha:
    """Ecuacion extraida de OMML (``m:oMath``) en un .docx."""

    latex: str
    html: str
    omml_sha256: str
    fuente: str = "omml"


@dataclass(frozen=True)
class FichaWordModelo:
    """Contenido de un .docx asociado a un modelo del catalogo."""

    archivo: str
    html: str
    imagenes: tuple[Path, ...]
    catalogo_id: str | None = None
    ruta: Path | None = None
    ecuaciones: tuple[EcuacionFicha, ...] = ()

def _normalizar(texto: str) -> str:
    t = unicodedata.normalize("NFKD", str(texto or ""))
    t = t.encode("ascii", "ignore").decode("ascii")
    return re.sub(r"[^a-z0-9]+", "", t.lower())


def _slug(texto: str) -> str:
    t = unicodedata.normalize("NFKD", str(texto or ""))
    t = t.encode("ascii", "ignore").decode("ascii")
    t = re.sub(r"[^A-Za-z0-9]+", "_", t).strip("_")
    return t or "ficha"


def _aliases_catalogo(entrada) -> set[str]:
    """Nombres normalizados con los que un archivo Word puede emparejarse."""
    from core.modelos.catalogo_impactos import titulo_modo_impacto
    from core.modelos.flujos import _ALIASES_FLUJO, nombre_esperado_diagrama

    aliases: set[str] = set()
    for bruto in (
        entrada.id,
        entrada.motor_id,
        entrada.motor_nombre,
        entrada.modo_fallo,
        titulo_modo_impacto(entrada),
        f"{entrada.familia} {entrada.modo_fallo}",
        f"{entrada.tipo_impacto} {entrada.modo_fallo}",
        nombre_esperado_diagrama(entrada.diagrama_modelo_id or entrada.motor_id),
    ):
        n = _normalizar(bruto)
        if n:
            aliases.add(n)

    mid = entrada.diagrama_modelo_id or entrada.motor_id
    for a in _ALIASES_FLUJO.get(mid, ()):
        n = _normalizar(a)
        if n:
            aliases.add(n)
    return aliases


def emparejar_nombre_ficha(nombre: str) -> str | None:
    """Devuelve ``entrada.id`` del catalogo o None.

    ``nombre`` puede ser hoja antigua, stem de archivo o titulo de modelo.
    """
    from core.modelos.catalogo_impactos import CATALOGO_MODOS_IMPACTO

    clave = _normalizar(Path(str(nombre or "")).stem)
    if not clave:
        return None

    mejor_id: str | None = None
    mejor_score = -1
    for entrada in CATALOGO_MODOS_IMPACTO:
        aliases = _aliases_catalogo(entrada)
        if clave in aliases:
            return entrada.id
        for a in aliases:
            if not a:
                continue
            if a in clave or clave in a:
                score = min(len(a), len(clave))
                if score > mejor_score:
                    mejor_score = score
                    mejor_id = entrada.id
    return mejor_id


def _listar_docx() -> list[Path]:
    if not CARPETA_FICHAS.is_dir():
        return []
    rutas: list[Path] = []
    for path in sorted(CARPETA_FICHAS.iterdir(), key=lambda p: p.name.lower()):
        if not path.is_file():
            continue
        if path.name.startswith("~$"):
            continue
        if path.suffix.lower() in _EXTENSIONES_WORD:
            rutas.append(path)
    return rutas


def _ext_imagen(data: bytes, content_type: str = "", nombre: str = "") -> str:
    ct = (content_type or "").lower()
    ext_nombre = Path(nombre).suffix.lower() if nombre else ""
    if ext_nombre in {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".emf", ".wmf"}:
        return ".jpg" if ext_nombre == ".jpeg" else ext_nombre
    if "jpeg" in ct or "jpg" in ct or data[:2] == b"\xff\xd8":
        return ".jpg"
    if "png" in ct or data[:8].startswith(b"\x89PNG"):
        return ".png"
    if "gif" in ct or data[:4] == b"GIF8":
        return ".gif"
    if "webp" in ct:
        return ".webp"
    return ".bin"


def _extraer_imagenes_media(ruta_docx: Path, media_dir: Path) -> tuple[Path, ...]:
    """Extrae imagenes embebidas del paquete OOXML a ``Fichas/_media/<slug>/``."""
    media_dir.mkdir(parents=True, exist_ok=True)
    guardadas: list[Path] = []
    try:
        with zipfile.ZipFile(ruta_docx) as zf:
            nombres = [
                n
                for n in zf.namelist()
                if n.startswith("word/media/") and not n.endswith("/")
            ]
            for i, nombre in enumerate(sorted(nombres)):
                data = zf.read(nombre)
                if not data:
                    continue
                ext = Path(nombre).suffix.lower() or _ext_imagen(data, nombre=nombre)
                if ext in {".emf", ".wmf"}:
                    # No mostrables directamente en Streamlit; se omiten.
                    continue
                dest = media_dir / f"img_{i}{ext}"
                if not dest.is_file() or dest.stat().st_size != len(data):
                    dest.write_bytes(data)
                guardadas.append(dest)
    except (OSError, zipfile.BadZipFile) as exc:
        _log.warning("No se pudieron extraer imagenes de %s: %s", ruta_docx.name, exc)
        return ()
    return tuple(guardadas)


def _extraer_imagenes_python_docx(ruta_docx: Path, media_dir: Path) -> tuple[Path, ...]:
    """Extrae imagenes via ``document.part.related_parts`` (python-docx)."""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        return ()

    media_dir.mkdir(parents=True, exist_ok=True)
    guardadas: list[Path] = []
    try:
        doc = Document(str(ruta_docx))
        parts = getattr(doc.part, "related_parts", {}) or {}
        i = 0
        for rel in doc.part.rels.values():
            try:
                if rel.reltype != RT.IMAGE:
                    continue
                part = rel.target_part
            except Exception:
                continue
            data = part.blob
            if not data:
                continue
            ct = getattr(part, "content_type", "") or ""
            nombre = getattr(part, "partname", None)
            nombre_s = str(nombre) if nombre else ""
            ext = _ext_imagen(data, content_type=ct, nombre=nombre_s)
            if ext in {".emf", ".wmf", ".bin"}:
                continue
            dest = media_dir / f"img_{i}{ext}"
            if not dest.is_file() or dest.stat().st_size != len(data):
                dest.write_bytes(data)
            guardadas.append(dest)
            i += 1
        # related_parts dict path (algunas builds)
        if not guardadas and parts:
            for part in parts.values():
                ct = getattr(part, "content_type", "") or ""
                if not str(ct).startswith("image/"):
                    continue
                data = getattr(part, "blob", None)
                if not data:
                    continue
                nombre_s = str(getattr(part, "partname", "") or "")
                ext = _ext_imagen(data, content_type=ct, nombre=nombre_s)
                if ext in {".emf", ".wmf", ".bin"}:
                    continue
                dest = media_dir / f"img_{i}{ext}"
                if not dest.is_file() or dest.stat().st_size != len(data):
                    dest.write_bytes(data)
                guardadas.append(dest)
                i += 1
    except Exception as exc:
        _log.warning("python-docx imagenes fallo en %s: %s", ruta_docx.name, exc)
        return ()
    return tuple(guardadas)


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

# Fallback solo si falla la extraccion OMML y el texto plano del doc coincide.
_FALLBACK_DELTA_RS_LATEX = (
    r"\Delta R_{S}=\max\left\{0,"
    r"\frac{T_{\mathrm{port}}+C_{\mathrm{conc}}}{365\cdot 24}"
    r"\cdot(HF_{s}-HF_{\mathrm{hist}})\right\}"
)
_FALLBACK_DELTA_RS_PLAIN = re.compile(
    r"\u0394?\s*R\s*S?\s*=\s*max|"
    r"T\s*port|C\s*conc|HF\s*s|HF\s*hist|365",
    re.IGNORECASE,
)


def _celda_a_html(texto: str) -> str:
    t = (texto or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if not t:
        return "&nbsp;"
    return html_lib.escape(t).replace("\n", "<br/>")


def _tc_texto(tc, tabla) -> str:
    """Texto de una celda OOXML ``w:tc`` via python-docx."""
    from docx.table import _Cell

    return _Cell(tc, tabla).text or ""


def _omml_local(tag: str | None) -> str:
    if not tag:
        return ""
    return tag.rsplit("}", 1)[-1]


def _omml_limpiar_texto(texto: str | None) -> str:
    """Quita ZWSP y basura tipografica de ``m:t``."""
    if not texto:
        return ""
    return (
        str(texto)
        .replace("\u200b", "")
        .replace("\u200c", "")
        .replace("\u200d", "")
        .replace("\ufeff", "")
    )


def _omml_hijos_utiles(el: ET.Element) -> list[ET.Element]:
    """Hijos OMML ignorando propiedades (``*Pr``) y ``ctrlPr``."""
    out: list[ET.Element] = []
    for child in list(el):
        name = _omml_local(child.tag)
        if not name or name.endswith("Pr") or name == "ctrlPr":
            continue
        out.append(child)
    return out


def _omml_primer_hijo(el: ET.Element, *nombres: str) -> ET.Element | None:
    wanted = set(nombres)
    for child in list(el):
        if _omml_local(child.tag) in wanted:
            return child
    return None


def _latex_escape_texto(texto: str) -> str:
    """Escapa caracteres LaTeX en texto literal OMML."""
    t = _omml_limpiar_texto(texto)
    if not t:
        return ""
    repl = {
        "\\": r"\backslash{}",
        "{": r"\{",
        "}": r"\}",
        "#": r"\#",
        "%": r"\%",
        "&": r"\&",
        "_": r"\_",
        "$": r"\$",
        "\u22c5": r"\cdot ",
        "\u00b7": r"\cdot ",
        "\u2212": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u0394": r"\Delta ",
        "\u03b4": r"\delta ",
    }
    out: list[str] = []
    for ch in t:
        out.append(repl.get(ch, ch))
    return "".join(out)

def _html_escape_math_texto(texto: str) -> str:
    t = (
        _omml_limpiar_texto(texto)
        .replace("\u22c5", "\u00b7")
        .replace("\u2212", "\u2212")
    )
    return html_lib.escape(t)


def _omml_a_latex(el: ET.Element | None) -> str:
    """Convierte un nodo OMML (``m:oMath`` u hijo) a LaTeX."""
    if el is None:
        return ""
    name = _omml_local(el.tag)

    if name in {"oMath", "oMathPara", "e", "num", "den", "sub", "sup", "fName"}:
        return "".join(_omml_a_latex(c) for c in _omml_hijos_utiles(el))

    if name == "r":
        partes: list[str] = []
        for child in list(el):
            if _omml_local(child.tag) == "t":
                partes.append(_latex_escape_texto(child.text))
        return "".join(partes)

    if name == "t":
        return _latex_escape_texto(el.text)

    if name == "sSub":
        base = _omml_a_latex(_omml_primer_hijo(el, "e"))
        sub = _omml_a_latex(_omml_primer_hijo(el, "sub"))
        return f"{{{base}}}_{{{sub}}}"

    if name == "sSup":
        base = _omml_a_latex(_omml_primer_hijo(el, "e"))
        sup = _omml_a_latex(_omml_primer_hijo(el, "sup"))
        return f"{{{base}}}^{{{sup}}}"

    if name == "sSubSup":
        base = _omml_a_latex(_omml_primer_hijo(el, "e"))
        sub = _omml_a_latex(_omml_primer_hijo(el, "sub"))
        sup = _omml_a_latex(_omml_primer_hijo(el, "sup"))
        return f"{{{base}}}_{{{sub}}}^{{{sup}}}"

    if name == "f":
        num = _omml_a_latex(_omml_primer_hijo(el, "num"))
        den = _omml_a_latex(_omml_primer_hijo(el, "den"))
        return rf"\frac{{{num}}}{{{den}}}"

    if name == "func":
        fname = _omml_a_latex(_omml_primer_hijo(el, "fName")).strip()
        arg = _omml_a_latex(_omml_primer_hijo(el, "e"))
        known = {
            "max": r"\max",
            "min": r"\min",
            "sin": r"\sin",
            "cos": r"\cos",
            "tan": r"\tan",
            "log": r"\log",
            "ln": r"\ln",
            "exp": r"\exp",
            "lim": r"\lim",
        }
        op = known.get(fname.lower(), rf"\mathrm{{{fname}}}" if fname else "")
        return f"{op}{arg}"

    if name == "d":
        d_pr = _omml_primer_hijo(el, "dPr")
        beg = "("
        end = ")"
        if d_pr is not None:
            beg_el = d_pr.find(f"{_M_NS}begChr")
            end_el = d_pr.find(f"{_M_NS}endChr")
            if beg_el is not None and beg_el.get(f"{_M_NS}val") is not None:
                beg = beg_el.get(f"{_M_NS}val") or beg
            if end_el is not None and end_el.get(f"{_M_NS}val") is not None:
                end = end_el.get(f"{_M_NS}val") or end
        inner = _omml_a_latex(_omml_primer_hijo(el, "e"))
        delim = {
            "{": (r"\left\{", r"\right\}"),
            "}": (r"\left\{", r"\right\}"),
            "[": (r"\left[", r"\right]"),
            "]": (r"\left[", r"\right]"),
            "(": (r"\left(", r"\right)"),
            ")": (r"\left(", r"\right)"),
            "|": (r"\left|", r"\right|"),
            "\u2016": (r"\left\|", r"\right\|"),
        }
        if beg in delim:
            left, right = delim[beg]
            return f"{left}{inner}{right}"
        left = _latex_escape_texto(beg)
        right = _latex_escape_texto(end)
        return f"{left}{inner}{right}"

    if name == "rad":
        deg = _omml_primer_hijo(el, "deg")
        base = _omml_a_latex(_omml_primer_hijo(el, "e"))
        if deg is not None and "".join(
            _omml_limpiar_texto(t.text) for t in deg.iter(f"{_M_NS}t")
        ):
            return rf"\sqrt[{_omml_a_latex(deg)}]{{{base}}}"
        return rf"\sqrt{{{base}}}"

    if name == "nary":
        nary_pr = _omml_primer_hijo(el, "naryPr")
        char = "\u2211"
        if nary_pr is not None:
            ch_el = nary_pr.find(f"{_M_NS}chr")
            if ch_el is not None and ch_el.get(f"{_M_NS}val"):
                char = ch_el.get(f"{_M_NS}val") or char
        op_map = {
            "\u2211": r"\sum",
            "\u220f": r"\prod",  # ?
            "\u222b": r"\int",  # ?
            "\u22c3": r"\bigcup",  # ?
        }
        op = op_map.get(char, rf"\mathrm{{{_latex_escape_texto(char)}}}")
        sub = _omml_a_latex(_omml_primer_hijo(el, "sub"))
        sup = _omml_a_latex(_omml_primer_hijo(el, "sup"))
        body = _omml_a_latex(_omml_primer_hijo(el, "e"))
        lim = ""
        if sub:
            lim += f"_{{{sub}}}"
        if sup:
            lim += f"^{{{sup}}}"
        return f"{op}{lim}{{{body}}}"

    # Nodos desconocidos: concatenar hijos utiles.
    return "".join(_omml_a_latex(c) for c in _omml_hijos_utiles(el))


def _omml_a_html(el: ET.Element | None) -> str:
    """HTML tipografico (fracciones/subindices) paralelo al LaTeX."""
    if el is None:
        return ""
    name = _omml_local(el.tag)

    if name in {"oMath", "oMathPara", "e", "num", "den", "sub", "sup", "fName"}:
        return "".join(_omml_a_html(c) for c in _omml_hijos_utiles(el))

    if name == "r":
        partes: list[str] = []
        for child in list(el):
            if _omml_local(child.tag) == "t":
                partes.append(_html_escape_math_texto(child.text))
        return "".join(partes)

    if name == "t":
        return _html_escape_math_texto(el.text)

    if name == "sSub":
        base = _omml_a_html(_omml_primer_hijo(el, "e"))
        sub = _omml_a_html(_omml_primer_hijo(el, "sub"))
        return f"{base}<sub>{sub}</sub>"

    if name == "sSup":
        base = _omml_a_html(_omml_primer_hijo(el, "e"))
        sup = _omml_a_html(_omml_primer_hijo(el, "sup"))
        return f"{base}<sup>{sup}</sup>"

    if name == "sSubSup":
        base = _omml_a_html(_omml_primer_hijo(el, "e"))
        sub = _omml_a_html(_omml_primer_hijo(el, "sub"))
        sup = _omml_a_html(_omml_primer_hijo(el, "sup"))
        return f"{base}<sub>{sub}</sub><sup>{sup}</sup>"

    if name == "f":
        num = _omml_a_html(_omml_primer_hijo(el, "num"))
        den = _omml_a_html(_omml_primer_hijo(el, "den"))
        return (
            f'<span class="frac"><span class="num">{num}</span>'
            f'<span class="den">{den}</span></span>'
        )

    if name == "func":
        fname = _omml_a_html(_omml_primer_hijo(el, "fName"))
        arg = _omml_a_html(_omml_primer_hijo(el, "e"))
        return f"<span>{fname}</span>{arg}"

    if name == "d":
        d_pr = _omml_primer_hijo(el, "dPr")
        beg = "("
        end = ")"
        if d_pr is not None:
            beg_el = d_pr.find(f"{_M_NS}begChr")
            end_el = d_pr.find(f"{_M_NS}endChr")
            if beg_el is not None and beg_el.get(f"{_M_NS}val") is not None:
                beg = beg_el.get(f"{_M_NS}val") or beg
            if end_el is not None and end_el.get(f"{_M_NS}val") is not None:
                end = end_el.get(f"{_M_NS}val") or end
        inner = _omml_a_html(_omml_primer_hijo(el, "e"))
        return f"{html_lib.escape(beg)}{inner}{html_lib.escape(end)}"

    if name == "rad":
        base = _omml_a_html(_omml_primer_hijo(el, "e"))
        return f"\u221a({base})"

    return "".join(_omml_a_html(c) for c in _omml_hijos_utiles(el))


def _omml_sha256(el: ET.Element) -> str:
    """Hash estable del XML OMML (lxml/oxml o ElementTree)."""
    raw: bytes
    xml_attr = getattr(el, "xml", None)
    if isinstance(xml_attr, str) and xml_attr:
        raw = xml_attr.encode("utf-8")
    else:
        try:
            raw = ET.tostring(el, encoding="utf-8")
        except Exception:
            raw = repr(el).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _ecuacion_desde_omath(om: ET.Element) -> EcuacionFicha | None:
    latex = (_omml_a_latex(om) or "").strip()
    html_eq = (_omml_a_html(om) or "").strip()
    if not latex and not html_eq:
        return None
    digest = _omml_sha256(om)
    return EcuacionFicha(
        latex=latex,
        html=html_eq,
        omml_sha256=digest,
        fuente="omml",
    )


def _extraer_omath_de_tc(tc) -> list[ET.Element]:
    return list(tc.findall(f".//{_M_NS}oMath"))


def _ecuacion_html_bloque(eq: EcuacionFicha) -> str:
    cuerpo = eq.html or html_lib.escape(eq.latex)
    return (
        f'<div class="pde-ficha-eq" data-omml-sha256="{html_lib.escape(eq.omml_sha256)}" '
        f'data-fuente="{html_lib.escape(eq.fuente)}">{cuerpo}</div>'
    )


def _tc_contenido_html(tc, tabla) -> tuple[str, tuple[EcuacionFicha, ...]]:
    """HTML de celda + ecuaciones OMML (texto economico + formula)."""
    texto = _tc_texto(tc, tabla)
    eqs: list[EcuacionFicha] = []
    for om in _extraer_omath_de_tc(tc):
        eq = _ecuacion_desde_omath(om)
        if eq is not None:
            eqs.append(eq)

    partes: list[str] = []
    if (texto or "").strip():
        partes.append(_celda_a_html(texto))
    for eq in eqs:
        partes.append(_ecuacion_html_bloque(eq))

    if not partes:
        return "&nbsp;", tuple(eqs)
    return "".join(partes), tuple(eqs)

def _tc_grid_span(tc) -> int:
    """``w:gridSpan`` (colspan); por defecto 1."""
    tc_pr = tc.find(f"{_W_NS}tcPr")
    if tc_pr is None:
        return 1
    gs = tc_pr.find(f"{_W_NS}gridSpan")
    if gs is None:
        return 1
    try:
        return max(1, int(gs.get(f"{_W_NS}val") or 1))
    except (TypeError, ValueError):
        return 1


def _tc_vmerge(tc) -> str | None:
    """Estado de fusion vertical OOXML.

    - ``None``: celda normal (sin vMerge)
    - ``\"restart\"``: inicio de fusion vertical
    - ``\"continue\"``: continuation (no emitir ``<td>``)

    En OOXML, ``w:vMerge`` sin ``w:val`` equivale a ``continue``.
    """
    tc_pr = tc.find(f"{_W_NS}tcPr")
    if tc_pr is None:
        return None
    vm = tc_pr.find(f"{_W_NS}vMerge")
    if vm is None:
        return None
    val = (vm.get(f"{_W_NS}val") or "").strip().lower()
    if not val or val == "continue":
        return "continue"
    return "restart"


@dataclass(frozen=True)
class _TcInfo:
    tc: object
    col: int
    colspan: int
    vmerge: str | None
    text: str
    html: str
    ecuaciones: tuple[EcuacionFicha, ...]


def _parsear_filas_ooxml(tabla) -> list[list[_TcInfo]]:
    """Lee filas/celdas reales del ``w:tbl`` (respeta vMerge/gridSpan)."""
    filas: list[list[_TcInfo]] = []
    for tr in tabla._tbl.findall(f"{_W_NS}tr"):
        col = 0
        celdas: list[_TcInfo] = []
        for tc in tr.findall(f"{_W_NS}tc"):
            colspan = _tc_grid_span(tc)
            html_celda, eqs = _tc_contenido_html(tc, tabla)
            celdas.append(
                _TcInfo(
                    tc=tc,
                    col=col,
                    colspan=colspan,
                    vmerge=_tc_vmerge(tc),
                    text=_tc_texto(tc, tabla),
                    html=html_celda,
                    ecuaciones=eqs,
                )
            )
            col += colspan
        filas.append(celdas)
    return filas


def _calcular_rowspan(filas: list[list[_TcInfo]], ri: int, celda: _TcInfo) -> int:
    """Numero de filas que ocupa un restart (o celda normal = 1)."""
    if celda.vmerge == "continue":
        return 0
    if celda.vmerge != "restart":
        return 1
    span = 1
    col = celda.col
    for rj in range(ri + 1, len(filas)):
        debajo = next((c for c in filas[rj] if c.col == col), None)
        if debajo is None or debajo.vmerge != "continue":
            break
        span += 1
    return span


def _tabla_a_html(tabla) -> tuple[str, tuple[EcuacionFicha, ...]]:
    """Convierte una tabla Word a HTML respetando fusiones OOXML.

    - ``w:vMerge`` restart -> ``rowspan`` con el texto una sola vez
    - ``w:vMerge`` continue (o sin val) -> no se emite ``<td>``
    - ``w:gridSpan`` -> ``colspan``
    - ``m:oMath`` en celda -> bloque HTML de ecuacion (+ LaTeX aparte)
    """
    filas = _parsear_filas_ooxml(tabla)
    filas_html: list[str] = []
    ecuaciones: list[EcuacionFicha] = []
    for ri, celdas in enumerate(filas):
        parts: list[str] = []
        tag = "th" if ri == 0 else "td"
        for celda in celdas:
            if celda.vmerge == "continue":
                continue
            rowspan = _calcular_rowspan(filas, ri, celda)
            attrs = ""
            if celda.colspan > 1:
                attrs += f' colspan="{celda.colspan}"'
            if rowspan > 1:
                attrs += f' rowspan="{rowspan}"'
            parts.append(f"<{tag}{attrs}>{celda.html}</{tag}>")
            ecuaciones.extend(celda.ecuaciones)
        filas_html.append("<tr>" + "".join(parts) + "</tr>")
    return "<table>" + "".join(filas_html) + "</table>", tuple(ecuaciones)


def _fallback_ecuacion_si_procede(ruta: Path, texto_doc: str) -> EcuacionFicha | None:
    """Fallback documentado solo si el .docx parece contener Delta Rs y no hay OMML."""
    stem = _normalizar(ruta.stem)
    if "francobordo" not in stem and "faltafrancobordo" not in stem:
        return None
    if not _FALLBACK_DELTA_RS_PLAIN.search(texto_doc or ""):
        return None
    latex = _FALLBACK_DELTA_RS_LATEX
    html_eq = (
        "\u0394R<sub>S</sub> = max{0, "
        '<span class="frac"><span class="num">T<sub>port</sub>+C<sub>conc</sub></span>'
        '<span class="den">365\u00b724</span></span>'
        "\u00b7(HF<sub>s</sub>\u2212HF<sub>hist</sub>)}"
    )
    digest = hashlib.sha256(f"fallback:{latex}".encode("utf-8")).hexdigest()
    _log.warning(
        "OMML no encontrado en %s; usando formula conocida Delta Rs "
        "(coincide con el texto del documento).",
        ruta.name,
    )
    return EcuacionFicha(
        latex=latex,
        html=html_eq,
        omml_sha256=digest,
        fuente="fallback_doc_match",
    )


def _html_desde_python_docx(ruta: Path) -> tuple[str, tuple[EcuacionFicha, ...]]:
    """Tablas (+ parrafos sueltos) a HTML con python-docx; extrae OMML."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(ruta))
    bloques: list[str] = []
    ecuaciones: list[EcuacionFicha] = []
    textos_planos: list[str] = []

    # Recorrer body en orden: parrafos y tablas intercalados.
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "tbl":
            tabla = Table(child, doc)
            html_tabla, eqs = _tabla_a_html(tabla)
            bloques.append(html_tabla)
            ecuaciones.extend(eqs)
            for row in tabla.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        textos_planos.append(t)
        elif tag == "p":
            para = Paragraph(child, doc)
            texto = (para.text or "").strip()
            if texto:
                textos_planos.append(texto)
                bloques.append(f"<p>{_celda_a_html(texto)}</p>")
            for om in child.findall(f".//{_M_NS}oMath"):
                eq = _ecuacion_desde_omath(om)
                if eq is not None:
                    ecuaciones.append(eq)
                    bloques.append(_ecuacion_html_bloque(eq))

    if not bloques and doc.tables:
        for tabla in doc.tables:
            html_tabla, eqs = _tabla_a_html(tabla)
            bloques.append(html_tabla)
            ecuaciones.extend(eqs)

    if not ecuaciones:
        fb = _fallback_ecuacion_si_procede(ruta, "\n".join(textos_planos))
        if fb is not None:
            ecuaciones.append(fb)
            # Insertar en el primer bloque de tabla si existe.
            if bloques and bloques[0].startswith("<table>"):
                bloques[0] = bloques[0].replace(
                    "</table>",
                    f"<tr><td colspan=\"3\">{_ecuacion_html_bloque(fb)}</td></tr></table>",
                    1,
                )
            else:
                bloques.append(_ecuacion_html_bloque(fb))

    if not bloques:
        return "", tuple(ecuaciones)

    body = "\n".join(bloques)
    html = f'<div class="pde-ficha-word"><style>{_CSS_FICHA}</style>{body}</div>'
    return html, tuple(ecuaciones)

def _html_desde_mammoth(ruta: Path) -> str:
    """Fallback opcional: convierte el .docx a HTML con mammoth."""
    import mammoth

    def _convert_image(image):
        with image.open() as image_bytes:
            raw = image_bytes.read()
        encoded = base64.b64encode(raw).decode("ascii")
        content_type = getattr(image, "content_type", None) or "image/png"
        return {"src": f"data:{content_type};base64,{encoded}"}

    with ruta.open("rb") as f:
        result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(_convert_image),
        )
    body = (result.value or "").strip()
    if not body:
        return ""
    return f'<div class="pde-ficha-word"><style>{_CSS_FICHA}</style>{body}</div>'


def _html_desde_docx(ruta: Path) -> tuple[str, tuple[EcuacionFicha, ...]]:
    """Convierte el .docx a HTML. Preferir python-docx; mammoth opcional."""
    errores: list[str] = []

    try:
        html, ecuaciones = _html_desde_python_docx(ruta)
        if html:
            return html, ecuaciones
        errores.append("python-docx: documento sin tablas/parrafos utiles")
    except ImportError as exc:
        errores.append(f"python-docx no instalado: {exc}")
    except Exception as exc:
        _log.exception("python-docx fallo al convertir %s", ruta.name)
        errores.append(f"python-docx: {type(exc).__name__}: {exc}")

    try:
        html = _html_desde_mammoth(ruta)
        if html:
            return html, ()
        errores.append("mammoth: conversion vacia")
    except ImportError:
        errores.append("mammoth no instalado (fallback omitido)")
    except Exception as exc:
        _log.exception("mammoth fallo al convertir %s", ruta.name)
        errores.append(f"mammoth: {type(exc).__name__}: {exc}")

    detalle = html_lib.escape(" | ".join(errores) if errores else "motivo desconocido")
    py_exe = html_lib.escape(sys.executable)
    _log.error(
        "No se pudo convertir %s [%s]: %s",
        ruta.name,
        sys.executable,
        " | ".join(errores),
    )
    return (
        (
            '<div class="pde-ficha-word"><p>'
            f"No se pudo convertir {html_lib.escape(ruta.name)}. "
            f"<code>{detalle}</code>"
            f" <small>(Python: <code>{py_exe}</code>)</small>"
            "</p></div>"
        ),
        (),
    )


def _imagenes_desde_docx(ruta: Path, media_dir: Path) -> tuple[Path, ...]:
    """Imagenes para el panel Esquema: python-docx, si no zip OOXML."""
    imgs = _extraer_imagenes_python_docx(ruta, media_dir)
    if imgs:
        return imgs
    return _extraer_imagenes_media(ruta, media_dir)


def _firma_carpeta_fichas() -> tuple[tuple[str, int, int], ...]:
    """Huella (name, size, mtime_ns) para detectar nuevas/removidas fichas."""
    items: list[tuple[str, int, int]] = []
    for ruta in _listar_docx():
        try:
            stat = ruta.stat()
        except OSError:
            continue
        items.append((ruta.name, int(stat.st_size), int(stat.st_mtime_ns)))
    return tuple(items)


@lru_cache(maxsize=32)
def _parsear_ficha(
    ruta_str: str,
    size: int,
    mtime_ns: int,
) -> tuple[str, tuple[Path, ...], tuple[EcuacionFicha, ...]]:
    """Parseo caro (HTML + imagenes + ecuaciones) cacheado por (ruta, size, mtime)."""
    del size, mtime_ns  # solo se usan como clave de cache
    ruta = Path(ruta_str)
    media = CARPETA_MEDIA / _slug(ruta.stem)
    imagenes = _imagenes_desde_docx(ruta, media)
    html, ecuaciones = _html_desde_docx(ruta)
    return html, imagenes, ecuaciones


def _cargar_fichas_word() -> dict[str, FichaWordModelo]:
    """Mapa catalogo_id -> ficha reconstruido en cada llamada.

    El parseo del .docx esta cacheado (por ruta+size+mtime), pero el
    emparejamiento con el catalogo se recalcula siempre para que los cambios
    en catalogo o en la carpeta ``Fichas/`` se apliquen sin reiniciar.
    """
    resultado: dict[str, FichaWordModelo] = {}
    for ruta in _listar_docx():
        try:
            stat = ruta.stat()
        except OSError:
            continue
        stem = ruta.stem
        catalogo_id = emparejar_nombre_ficha(stem)
        html, imagenes, ecuaciones = _parsear_ficha(
            str(ruta), int(stat.st_size), int(stat.st_mtime_ns)
        )
        key = catalogo_id or f"docx:{stem}"
        resultado[key] = FichaWordModelo(
            archivo=ruta.name,
            html=html,
            imagenes=imagenes,
            catalogo_id=catalogo_id,
            ruta=ruta,
            ecuaciones=ecuaciones,
        )
    return resultado


def invalidar_cache_fichas() -> None:
    _parsear_ficha.cache_clear()


def ficha_word_por_catalogo_id(catalogo_id: str) -> FichaWordModelo | None:
    if not (catalogo_id or "").strip():
        return None
    return _cargar_fichas_word().get(catalogo_id.strip())


def ficha_word_por_entrada(entrada) -> FichaWordModelo | None:
    """Resuelve ficha Word para una entrada del catalogo."""
    if entrada is None:
        return None
    directa = ficha_word_por_catalogo_id(entrada.id)
    if directa is not None:
        return directa
    for ficha in _cargar_fichas_word().values():
        if ficha.catalogo_id == entrada.id:
            return ficha
        if emparejar_nombre_ficha(ficha.archivo) == entrada.id:
            return ficha
    return None


def imagenes_ficha_por_entrada(entrada) -> tuple[Path, ...]:
    ficha = ficha_word_por_entrada(entrada)
    if ficha is None:
        return ()
    return ficha.imagenes


def fichas_disponibles() -> tuple[str, ...]:
    """Nombres de archivos Word detectados en ``Fichas/``."""
    return tuple(p.name for p in _listar_docx())
